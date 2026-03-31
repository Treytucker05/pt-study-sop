import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

def set_table_borders(table):
    # Find or create tblPr
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
        
    # Remove existing tblBorders if present to replace them
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
        
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '18') # 18 eighths of a point = 2.25pt width (very visible)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def format_document():
    input_path = r'C:\Users\treyt\Downloads\HTI_Balance_Checkoff_OnePage_Table_v2.docx'
    output_path = r'C:\Users\treyt\Downloads\HTI_Balance_Checkoff_OnePage_Table_v3.docx'
    
    doc = docx.Document(input_path)
    
    # 1. Adjust margins to be as tight as reasonably possible
    for section in doc.sections:
        section.top_margin = Inches(0.15)
        section.bottom_margin = Inches(0.15)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
        
    # 2. Format tables
    for table in doc.tables:
        set_table_borders(table)
        
        # Optionally, Autofit behavior
        table.autofit = True
        
        for row in table.rows:
            # Reduce row height if set explicitly
            row.height = None
            for cell in row.cells:
                # Optional: adjust cell margins if necessary, but standard docx api doesn't easily expose cell padding
                for paragraph in cell.paragraphs:
                    # Remove spacing before/after and set single line spacing
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    
                    for run in paragraph.runs:
                        # Reduce font size slightly to ensure fitting
                        if run.font.size is not None:
                            run.font.size = Pt(max(8, run.font.size.pt - 1.5))
                        else:
                            run.font.size = Pt(9.5)

    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

if __name__ == '__main__':
    format_document()
