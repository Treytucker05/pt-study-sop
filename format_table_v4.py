import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

def set_table_borders(table):
    # Find or create tblPr
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
        
    # Remove existing tblBorders if present to replace them
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
        
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8') # 8 eighths of a point = 1.0pt width (visible but not too thick)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

bullets_data = {
    1: ["State condition: Eyes closed on foam", "Goal: Sit upright for 5 mins at baseball game", "Position: Sitting, static balance, reactive motor strategy"],
    2: ["State patient is seated safely on foam", "Firm surface, feet supported", "Guarding closely for safety"],
    3: ["Assessing patient in assigned sensory condition", "Purpose: Examine balance without vision & less reliable somatosensory input"],
    4: ["Instruction: 'Sit upright, feet supported, eyes closed'", "Instruction: 'Maintain balance until I say stop'", "1 timed trial (up to 30 sec)", "Stop if: eyes open, loss of balance, needs manual assist, or excessive UE support"],
    5: ["Vision removed, somatosensory less reliable (foam)", "Patient must rely on vestibular input & postural control", "Regression: Open eyes (increases visual weighting, easier)"],
    6: ["Intervention: Seated upright balance", "On foam, eyes closed", "Small unpredictable multidirectional perturbations"],
    7: ["Matches sitting: Patient stays seated", "Matches static: Holds upright posture", "Matches reactive: Responds to unexpected perturbations", "Matches goal: Baseball game has unpredictable movement & distraction"],
    8: ["Instruction: 'Sit upright, eyes closed'", "Instruction: 'Recover balance each time I nudge you'"],
    9: ["Mode: Motor coord & balance training", "Intensity: Challenging but successful", "Duration: 20-30 sec bout (4-8 perturbations)", "Frequency: 2-3 rounds total in session"],
    10: ["Correct intensity: Shows corrective responses, recovers safely", "Too hard: Repeatedly opens eyes, grabs for support, needs manual assist", "Too easy: Minimal corrective response, no meaningful challenge"],
    11: ["Choose Progression OR Regression:", "Progression: Completed safely with minimal loss of control. Expect to maintain sitting with more challenge.", "Regression: Repeated loss of control, required lower challenge. Expect to maintain sitting with less challenge."],
    12: ["Instruction: 'Now changing the task, maintain upright sitting'", "Perform updated intervention (20-30 sec bout)", "Addresses goal (trains seated postural control without vision/somatosensory)", "Concludes checkoff"]
}

def format_document():
    input_path = r'C:\Users\treyt\Downloads\HTI_Balance_Checkoff_OnePage_Table_v2.docx'
    output_path = r'C:\Users\treyt\Downloads\HTI_Balance_Checkoff_OnePage_Table_v4.docx'
    
    doc = docx.Document(input_path)
    
    # 1. Adjust margins to be as tight as reasonably possible
    for section in doc.sections:
        section.top_margin = Inches(0.15)
        section.bottom_margin = Inches(0.15)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
        
    # 2. Format tables
    table = doc.tables[0]
    set_table_borders(table)
    table.autofit = True
    
    for i, row in enumerate(table.rows):
        row.height = None # clear explicit heights
        
        # Rewrite the "What to say" column if it's not the header row
        if i > 0 and i in bullets_data:
            cell = row.cells[2]
            cell.text = ""
            for bullet in bullets_data[i]:
                p = cell.add_paragraph(f"• {bullet}")
                p.style = doc.styles['Normal']
            
            # Clean up any empty first paragraph left over by cell.text = ""
            if len(cell.paragraphs) > 1 and cell.paragraphs[0].text.strip() == "":
                p = cell.paragraphs[0]._element
                p.getparent().remove(p)

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Remove spacing before/after and set single line spacing
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                
                for run in paragraph.runs:
                    # Reduce font size to ensure fitting and keep it glanceable
                    run.font.size = Pt(9.5)
                    # Let's bold the bullets in column 1 and 2 slightly if they have prefixes, but for now just standard text
                    
    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

if __name__ == '__main__':
    format_document()
