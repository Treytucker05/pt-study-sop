"""Semester intake routes for starting a fresh term inside Study."""

from __future__ import annotations

import os
import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any

from flask import Blueprint, jsonify, request

from course_wheel_sync import ensure_course_in_wheel
from dashboard.api_tutor_materials import _launch_materials_sync_job
from db_setup import get_connection

semester_intake_bp = Blueprint(
    "semester_intake", __name__, url_prefix="/api/semester-intake"
)

SUPPORTED_EXTS = {".docx", ".markdown", ".md", ".mp4", ".pdf", ".ppt", ".pptx", ".txt"}
ADMIN_KEYWORDS = {
    "background",
    "bls",
    "cpi",
    "cpr",
    "drug",
    "exxat",
    "immunization",
    "naloxone",
    "resume",
    "tdap",
    "titer",
}
MONTHS = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}


def _resolve_root(raw: Any) -> Path:
    value = (
        raw
        or os.environ.get("TUTOR_MATERIALS_DIR")
        or os.environ.get("PT_SCHOOL_MATERIALS_DIR")
        or ""
    )
    folder = str(value).strip().strip('"').strip("'")
    if not folder:
        raise ValueError("folder_path is required")
    root = Path(folder).expanduser()
    if not root.exists() or not root.is_dir():
        raise FileNotFoundError(f"Folder not found: {folder}")
    return root.resolve()


def _normalize_rel_path(raw_path: Any) -> str:
    raw = str(raw_path or "").strip().replace("\\", "/")
    if not raw or raw.startswith("/") or re.match(r"^[A-Za-z]:", raw):
        raise ValueError("selected file paths must be relative")
    parts = [part for part in raw.split("/") if part and part != "."]
    if not parts or any(part == ".." for part in parts):
        raise ValueError("selected file paths cannot traverse parent folders")
    return "/".join(parts)


def _validate_rel_file(root: Path, rel_path: str) -> Path:
    normalized = _normalize_rel_path(rel_path)
    candidate = (root / normalized).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"File is outside the intake folder: {rel_path}") from exc
    if not candidate.exists() or not candidate.is_file():
        raise FileNotFoundError(f"Selected file not found: {rel_path}")
    if candidate.suffix.lower() not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type selected: {rel_path}")
    return candidate


def _extract_setup_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".docx":
        import docx

        doc = docx.Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    if ext == ".pdf":
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for table in page.extract_tables():
                    for row in table:
                        cells = [str(cell).strip() if cell else "" for cell in row]
                        if any(cells):
                            text += "\n" + " | ".join(cells)
                if text.strip():
                    pages.append(text)
        return "\n\n".join(pages)
    return ""


def _coerce_year(raw_year: str | None, fallback_year: int) -> int:
    if not raw_year:
        return fallback_year
    year = int(raw_year)
    if year < 100:
        return 2000 + year
    return year


def _iso_date(month: int, day: int, year: int) -> str | None:
    try:
        return datetime(year, month, day).date().isoformat()
    except ValueError:
        return None


def _first_date(segment: str, *, fallback_year: int) -> str | None:
    slash = re.search(
        r"\b(?P<month>\d{1,2})/(?P<day>\d{1,2})(?:/(?P<year>\d{2,4}))?\b",
        segment,
    )
    if slash:
        return _iso_date(
            int(slash.group("month")),
            int(slash.group("day")),
            _coerce_year(slash.group("year"), fallback_year),
        )

    month_name = re.search(
        r"\b(?P<month>Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t|tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?\s+(?P<day>\d{1,2})(?:,\s*(?P<year>\d{2,4}))?\b",
        segment,
        flags=re.IGNORECASE,
    )
    if month_name:
        month = MONTHS[month_name.group("month").rstrip(".").lower()]
        return _iso_date(
            month,
            int(month_name.group("day")),
            _coerce_year(month_name.group("year"), fallback_year),
        )
    return None


def _extract_time_range(segment: str) -> tuple[str | None, str | None]:
    match = re.search(
        r"\b(?P<start>\d{1,2}(?::\d{2})?)\s*(?:-|to|–|—)\s*(?P<end>\d{1,2}(?::\d{2})?)\s*(?P<ampm>a\.?m\.?|p\.?m\.?)\b",
        segment,
        flags=re.IGNORECASE,
    )
    if not match:
        return None, None
    suffix = match.group("ampm").replace(".", "").lower()
    return f"{match.group('start')} {suffix}", f"{match.group('end')} {suffix}"


def _clean_event_title(segment: str) -> str:
    cleaned = re.sub(r"\s+", " ", segment).strip(" -|")
    cleaned = re.sub(r"\bBEFORE\b.*$", "", cleaned, flags=re.IGNORECASE).strip()
    return cleaned[:180] or "Course event"


def _event_type_for_segment(segment: str) -> str:
    lowered = segment.lower()
    if "quiz" in lowered:
        return "quiz"
    if "exam" in lowered:
        return "exam"
    if "due" in lowered or "post" in lowered or "reflection" in lowered:
        return "assignment"
    if "class" in lowered or "lecture" in lowered or "lab" in lowered:
        return "lecture"
    return "other"


def _parse_setup_modules(text: str) -> list[dict[str, Any]]:
    modules: list[dict[str, Any]] = []
    seen: set[str] = set()
    for match in re.finditer(
        r"\bWEEK\s+(?P<week>\d+)\s*:\s*(?P<week_label>[^\n|]+)(?P<body>.*?)(?=\n\s*WEEK\s+\d+\s*:|\Z)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    ):
        week = int(match.group("week"))
        body = match.group("body")[:1200]
        module_match = re.search(
            r"\bModule\s+(?P<num>\d+)\s*:\s*(?P<title>[^\n|]+)",
            body,
            flags=re.IGNORECASE,
        )
        title = (
            module_match.group("title").strip()
            if module_match
            else match.group("week_label").strip()
        )
        title = re.sub(r"\s+", " ", title).strip(" -|")
        if not title:
            continue
        name = f"Week {week}: {title}"
        key = name.lower()
        if key in seen:
            continue
        seen.add(key)
        modules.append({"name": name, "orderIndex": week, "objectives": []})
    return modules


def _parse_setup_events(text: str, *, fallback_year: int) -> list[dict[str, Any]]:
    events: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    segments: list[str] = []
    for line in text.splitlines():
        for segment in re.split(r"\s+\|\s+", line):
            cleaned = re.sub(r"\s+", " ", segment).strip()
            if cleaned:
                segments.append(cleaned)

    for segment in segments:
        if re.match(r"^WEEK\s+\d+\s*:", segment, flags=re.IGNORECASE):
            continue
        date_val = _first_date(segment, fallback_year=fallback_year)
        if not date_val:
            continue
        lowered = segment.lower()
        if not any(
            marker in lowered
            for marker in ("class", "lecture", "lab", "due", "quiz", "exam", "post", "reflection")
        ):
            continue
        start_time, end_time = _extract_time_range(segment)
        event_type = _event_type_for_segment(segment)
        title = _clean_event_title(segment)
        key = (event_type, title.lower(), date_val)
        if key in seen:
            continue
        seen.add(key)
        events.append(
            {
                "type": event_type,
                "title": title,
                "date": date_val,
                "dueDate": date_val if event_type in {"assignment", "quiz", "exam"} else None,
                "startTime": start_time,
                "endTime": end_time,
                "notes": segment,
            }
        )
    return events


def _draft_setup_from_files(root: Path, item: dict[str, Any]) -> dict[str, Any]:
    selected_setup_files: list[str] = []
    for field_name in ("syllabus_files", "schedule_files"):
        for raw_path in item.get(field_name) or []:
            rel_path = _normalize_rel_path(raw_path)
            if rel_path not in selected_setup_files:
                selected_setup_files.append(rel_path)

    modules: list[dict[str, Any]] = []
    events: list[dict[str, Any]] = []
    errors: list[str] = []
    parsed_files = 0
    fallback_year = datetime.now().year

    for rel_path in selected_setup_files:
        try:
            file_path = _validate_rel_file(root, rel_path)
            text = _extract_setup_text(file_path)
            if not text.strip():
                continue
            parsed_files += 1
            modules.extend(_parse_setup_modules(text))
            events.extend(_parse_setup_events(text, fallback_year=fallback_year))
        except Exception as exc:
            errors.append(f"{rel_path}: {exc}")

    return {
        "modules": modules,
        "events": events,
        "parsedFiles": parsed_files,
        "errors": errors,
    }


def _clean_course_name(folder_name: str) -> str:
    cleaned = re.sub(r"^\d+[_\-\s]+", "", folder_name).replace("_", " ").strip()
    return re.sub(r"\s+", " ", cleaned) or folder_name


def _is_global_schedule_folder(name: str) -> bool:
    lowered = name.lower()
    return "schedule" in lowered and (lowered.startswith("00") or "class" in lowered)


def _is_admin_path(rel_path: str) -> bool:
    lowered = rel_path.lower()
    parts = lowered.split("/")
    if parts and (parts[0].startswith(".") or parts[0].startswith("90")):
        return True
    return any(keyword in lowered for keyword in ADMIN_KEYWORDS)


def _roles_for_path(rel_path: str) -> list[str]:
    lowered = rel_path.lower()
    filename = Path(rel_path).name.lower()
    parts = Path(rel_path).parts
    roles: list[str] = []
    if "syllabus" in lowered or "course map" in lowered:
        roles.append("syllabus")
    if "schedule" in lowered or "calendar" in lowered or "course map" in lowered:
        roles.append("schedule")
    if (
        not roles
        and len(parts) == 2
        and re.search(r"\bphyt[\s_-]*\d{4}(?!\d)", filename, flags=re.IGNORECASE)
    ):
        roles.append("syllabus")
    if not roles:
        roles.append("material")
    return roles


def _file_payload(root: Path, rel_path: str, roles: list[str]) -> dict[str, Any]:
    file_path = root / rel_path
    stat = file_path.stat()
    return {
        "path": rel_path,
        "name": file_path.name,
        "roles": roles,
        "size": int(stat.st_size),
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


def _read_existing_course_summary() -> dict[str, int]:
    conn = get_connection()
    try:
        rows = conn.execute("SELECT id, name, code FROM courses").fetchall()
    finally:
        conn.close()
    summary: dict[str, int] = {}
    for row in rows:
        course_id = int(row[0])
        for value in (row[1], row[2]):
            key = str(value or "").strip().lower()
            if key:
                summary[key] = course_id
    return summary


def _classify_intake_folder(root: Path) -> dict[str, Any]:
    existing_courses = _read_existing_course_summary()
    course_map: dict[str, dict[str, Any]] = {}
    global_schedule_files: list[dict[str, Any]] = []
    ignored_files: list[dict[str, Any]] = []
    unassigned_material_files: list[dict[str, Any]] = []

    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = sorted(
            [name for name in dirnames if not name.startswith(".")],
            key=str.lower,
        )
        for filename in sorted(filenames, key=str.lower):
            file_path = Path(dirpath) / filename
            if file_path.suffix.lower() not in SUPPORTED_EXTS:
                continue
            rel_path = os.path.relpath(file_path, root).replace("\\", "/")
            top_folder = rel_path.split("/", 1)[0]
            roles = _roles_for_path(rel_path)
            payload = _file_payload(root, rel_path, roles)

            if _is_admin_path(rel_path):
                ignored_files.append(payload)
                continue

            if _is_global_schedule_folder(top_folder):
                if "schedule" in roles:
                    global_schedule_files.append(payload)
                else:
                    unassigned_material_files.append(payload)
                continue

            course = course_map.setdefault(
                top_folder,
                {
                    "name": _clean_course_name(top_folder),
                    "folder_path": top_folder,
                    "syllabus_files": [],
                    "schedule_files": [],
                    "material_files": [],
                },
            )
            if "syllabus" in roles:
                course["syllabus_files"].append(payload)
            if "schedule" in roles:
                course["schedule_files"].append(payload)
            if roles == ["material"]:
                course["material_files"].append(payload)

    courses = []
    syllabus_count = 0
    schedule_count = len(global_schedule_files)
    material_count = len(unassigned_material_files)
    for course in sorted(course_map.values(), key=lambda item: item["folder_path"].lower()):
        course_name_key = str(course["name"]).lower()
        course_id = existing_courses.get(course_name_key)
        syllabus_count += len(course["syllabus_files"])
        schedule_count += len(course["schedule_files"])
        material_count += len(course["material_files"])
        course["readiness"] = {
            "course": "exists" if course_id else "missing",
            "syllabus": "found" if course["syllabus_files"] else "missing",
            "schedule": "found" if course["schedule_files"] else "missing",
            "materials": "found" if course["material_files"] else "missing",
            "embeddings": "pending",
            "readyForTutor": False,
        }
        if course_id:
            course["course_id"] = course_id
        courses.append(course)

    return {
        "ok": True,
        "folder": str(root),
        "courses": courses,
        "global_schedule_files": global_schedule_files,
        "unassigned_material_files": unassigned_material_files,
        "ignored_files": ignored_files,
        "counts": {
            "courses": len(courses),
            "syllabus_files": syllabus_count,
            "schedule_files": schedule_count,
            "material_files": material_count,
            "ignored_files": len(ignored_files),
        },
    }


def _find_or_create_course(
    cur: sqlite3.Cursor,
    *,
    name: str,
    code: str | None,
    term: str | None,
) -> tuple[int, bool]:
    if code:
        cur.execute(
            "SELECT id FROM courses WHERE lower(COALESCE(code, '')) = lower(?) LIMIT 1",
            (code,),
        )
        row = cur.fetchone()
        if row:
            return int(row[0]), False

    cur.execute("SELECT id FROM courses WHERE lower(name) = lower(?) LIMIT 1", (name,))
    row = cur.fetchone()
    if row:
        if code:
            cur.execute("UPDATE courses SET code = COALESCE(code, ?) WHERE id = ?", (code, row[0]))
        return int(row[0]), False

    now = datetime.now().isoformat()
    cur.execute(
        """
        INSERT INTO courses (name, code, term, default_study_mode, created_at)
        VALUES (?, ?, ?, 'Core', ?)
        """,
        (name, code, term, now),
    )
    return int(cur.lastrowid), True


def _insert_modules_and_objectives(
    cur: sqlite3.Cursor,
    *,
    course_id: int,
    modules: list[dict[str, Any]],
) -> tuple[int, int]:
    modules_created = 0
    objectives_created = 0
    now = datetime.now().isoformat()

    for index, module in enumerate(modules):
        name = str(module.get("name") or "").strip()
        if not name:
            continue
        order_index = int(module.get("orderIndex", index) or 0)
        cur.execute(
            "SELECT id FROM modules WHERE course_id = ? AND lower(name) = lower(?) LIMIT 1",
            (course_id, name),
        )
        row = cur.fetchone()
        if row:
            module_id = int(row[0])
        else:
            cur.execute(
                """
                INSERT INTO modules (
                    course_id, name, order_index, files_downloaded,
                    notebooklm_loaded, sources, created_at, updated_at
                ) VALUES (?, ?, ?, 0, 0, NULL, ?, ?)
                """,
                (course_id, name, order_index, now, now),
            )
            module_id = int(cur.lastrowid)
            modules_created += 1

        for objective in module.get("objectives") or []:
            title = str(objective.get("title") or "").strip()
            if not title:
                continue
            lo_code = str(objective.get("loCode") or "").strip() or None
            cur.execute(
                """
                SELECT id FROM learning_objectives
                WHERE course_id = ?
                  AND COALESCE(module_id, 0) = ?
                  AND lower(title) = lower(?)
                LIMIT 1
                """,
                (course_id, module_id, title),
            )
            if cur.fetchone():
                continue
            cur.execute(
                """
                INSERT INTO learning_objectives (
                    course_id, module_id, lo_code, title, status, created_at, updated_at
                ) VALUES (?, ?, ?, ?, 'not_started', ?, ?)
                """,
                (course_id, module_id, lo_code, title, now, now),
            )
            objectives_created += 1

    return modules_created, objectives_created


def _insert_schedule_events(
    cur: sqlite3.Cursor,
    *,
    course_id: int,
    events: list[dict[str, Any]],
) -> int:
    created = 0
    now = datetime.now().isoformat()
    for event in events:
        event_type = str(event.get("type") or "other").strip() or "other"
        title = str(event.get("title") or "").strip()
        if not title:
            continue
        date_val = event.get("date") or event.get("dueDate")
        due_date = event.get("dueDate")
        start_time = event.get("startTime")
        end_time = event.get("endTime")
        notes = event.get("notes")
        cur.execute(
            """
            SELECT id FROM course_events
            WHERE course_id = ? AND type = ? AND title = ?
              AND COALESCE(date, '') = COALESCE(?, '')
              AND COALESCE(due_date, '') = COALESCE(?, '')
            LIMIT 1
            """,
            (course_id, event_type, title, date_val, due_date),
        )
        if cur.fetchone():
            continue
        cur.execute(
            """
            INSERT INTO course_events (
                course_id, type, title, date, due_date, time, end_time,
                raw_text, status, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?, ?)
            """,
            (
                course_id,
                event_type,
                title,
                date_val,
                due_date,
                start_time,
                end_time,
                notes,
                now,
                now,
            ),
        )
        created += 1
    return created


@semester_intake_bp.route("/preview", methods=["POST"])
def preview_semester_intake():
    data = request.get_json(silent=True) or {}
    try:
        root = _resolve_root(data.get("folder_path"))
        return jsonify(_classify_intake_folder(root)), 200
    except (ValueError, FileNotFoundError) as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@semester_intake_bp.route("/apply", methods=["POST"])
def apply_semester_intake():
    data = request.get_json(silent=True) or {}
    try:
        root = _resolve_root(data.get("folder_path"))
        courses_payload = data.get("courses") or []
        if not isinstance(courses_payload, list):
            return jsonify({"error": "courses must be an array"}), 400

        courses_created = 0
        courses_updated = 0
        modules_created = 0
        objectives_created = 0
        events_created = 0
        setup_files_parsed = 0
        setup_parse_errors: list[str] = []
        material_sync_jobs: list[dict[str, Any]] = []
        pending_material_syncs: list[dict[str, Any]] = []
        applied_courses: list[dict[str, Any]] = []

        conn = get_connection()
        try:
            cur = conn.cursor()
            for item in courses_payload:
                if not isinstance(item, dict):
                    continue
                name = str(item.get("name") or "").strip()
                if not name:
                    name = _clean_course_name(str(item.get("folder_path") or "Course"))
                code = str(item.get("code") or "").strip() or None
                term = str(item.get("term") or "").strip() or None
                course_id, created = _find_or_create_course(
                    cur,
                    name=name,
                    code=code,
                    term=term,
                )
                if created:
                    courses_created += 1
                else:
                    courses_updated += 1
                applied_courses.append({"id": course_id, "name": name, "code": code})

                setup_draft = _draft_setup_from_files(root, item)
                setup_files_parsed += int(setup_draft["parsedFiles"])
                setup_parse_errors.extend(setup_draft["errors"])

                syllabus = item.get("syllabus") if isinstance(item.get("syllabus"), dict) else {}
                syllabus_modules = [
                    *(syllabus.get("modules") or []),
                    *setup_draft["modules"],
                ]
                created_modules, created_objectives = _insert_modules_and_objectives(
                    cur,
                    course_id=course_id,
                    modules=syllabus_modules,
                )
                modules_created += created_modules
                objectives_created += created_objectives

                schedule = item.get("schedule") if isinstance(item.get("schedule"), dict) else {}
                schedule_events = [
                    *(schedule.get("events") or []),
                    *setup_draft["events"],
                ]
                events_created += _insert_schedule_events(
                    cur,
                    course_id=course_id,
                    events=schedule_events,
                )

                material_files = [
                    _normalize_rel_path(path)
                    for path in item.get("material_files") or []
                ]
                for rel_path in material_files:
                    _validate_rel_file(root, rel_path)
                if material_files:
                    pending_material_syncs.append(
                        {"courseId": course_id, "selectedFiles": set(material_files)}
                    )

            conn.commit()
        finally:
            conn.close()

        for course in applied_courses:
            ensure_course_in_wheel(int(course["id"]), name=str(course["name"]))
        for sync in pending_material_syncs:
            job_id = _launch_materials_sync_job(
                root,
                None,
                selected_files=sync["selectedFiles"],
                course_id=int(sync["courseId"]),
            )
            material_sync_jobs.append({"courseId": int(sync["courseId"]), "jobId": job_id})

        return jsonify(
            {
                "ok": True,
                "coursesCreated": courses_created,
                "coursesUpdated": courses_updated,
                "modulesCreated": modules_created,
                "objectivesCreated": objectives_created,
                "eventsCreated": events_created,
                "setupFilesParsed": setup_files_parsed,
                "setupParseErrors": setup_parse_errors,
                "materialSyncJobs": material_sync_jobs,
                "courses": applied_courses,
            }
        ), 200
    except (ValueError, FileNotFoundError) as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500
